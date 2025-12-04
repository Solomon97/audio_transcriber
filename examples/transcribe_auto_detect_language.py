#!/usr/bin/env python3
"""
transcribe_auto_detect.py — Example script for multi-language audio transcription.

This script automatically detects the language of your audio and processes it
accordingly. It supports English, Chinese, Japanese, Korean, and many other
languages that Whisper can transcribe.

Workflow:
    1. Transcribe audio using faster-whisper (auto-detects language)
    2. Export the raw transcript to a Word document (.docx)
    3. Clean up the transcript using a language-appropriate LLM
    4. Summarize the transcript using an LLM
    5. Export both cleaned and summarized versions to Word documents

Prerequisites:
    - Install the package: pip install -e .[ollama]
    - Have Ollama running locally (default: http://localhost:11434)
    - Pull models for your target languages:
        ollama pull yi:9b          # Chinese
        ollama pull llama3.1:8b    # English
        ollama pull qwen2.5:7b     # Japanese, Korean, general fallback

Usage:
    python transcribe_auto_detect.py <audio_file>

Examples:
    python transcribe_auto_detect.py meeting.wav
    python transcribe_auto_detect.py podcast.mp3
    python transcribe_auto_detect.py interview.m4a

Output:
    Creates three .docx files in the same directory as the input audio:
        - <filename>_raw.docx      : Raw transcript with timestamps
        - <filename>_cleaned.docx  : LLM-cleaned transcript
        - <filename>_summary.docx  : LLM-generated summary

Language-Model Mapping:
    The script automatically selects the best LLM for each language:
        - Chinese (zh)  → yi:9b
        - English (en)  → llama3.1:8b
        - Japanese (ja) → qwen2.5:7b
        - Korean (ko)   → qwen2.5:7b
        - Other         → qwen2.5:7b (default)

Notes:
    - For best accuracy, the script uses the 'large-v3' Whisper model
    - Language is auto-detected from the first 30 seconds of audio
    - You can override detection by modifying the transcribe() call
"""

import sys
from pathlib import Path

from docx import Document

from audio_transcriber import Transcriber, PostProcessor


# =============================================================================
# Configuration
# =============================================================================

# Human-readable language names for display
LANGUAGE_NAMES: dict[str, str] = {
    "zh": "Chinese",
    "en": "English",
    "ja": "Japanese",
    "ko": "Korean",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "ru": "Russian",
    "pt": "Portuguese",
    "it": "Italian",
    "nl": "Dutch",
    "pl": "Polish",
    "vi": "Vietnamese",
    "th": "Thai",
    "id": "Indonesian",
    "ar": "Arabic",
    "hi": "Hindi",
}


# =============================================================================
# Helper Functions
# =============================================================================


def get_language_name(code: str) -> str:
    """
    Get the human-readable name for a language code.

    Args:
        code: ISO 639-1 language code (e.g., 'zh', 'en', 'ja').

    Returns:
        Human-readable language name, or the code itself if unknown.
    """
    return LANGUAGE_NAMES.get(code, code.upper())


def save_text_to_docx(
    text: str,
    output_path: Path,
    title: str,
    language: str,
    duration: float,
) -> Path:
    """
    Save text content to a Word document with metadata.

    Args:
        text: The text content to save.
        output_path: Path where the .docx file will be saved.
        title: Document title/heading.
        language: Detected language code (e.g., 'zh', 'en').
        duration: Audio duration in seconds.

    Returns:
        Path to the created document.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    # Add metadata section
    language_display = f"{get_language_name(language)} ({language})"
    doc.add_paragraph(f"Language: {language_display}")
    doc.add_paragraph(f"Duration: {format_duration(duration)}")
    doc.add_paragraph()  # Empty line for spacing

    # Add main content
    doc.add_heading("Content", level=2)
    doc.add_paragraph(text)

    doc.save(output_path)
    return output_path


def format_duration(seconds: float) -> str:
    """
    Format seconds into a human-readable duration string.

    Args:
        seconds: Duration in seconds.

    Returns:
        Formatted string like "5m 30s" or "1h 15m 30s".
    """
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)

    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    elif minutes > 0:
        return f"{minutes}m {secs}s"
    else:
        return f"{secs}s"


# =============================================================================
# Main Processing Pipeline
# =============================================================================


def process_audio(audio_file: str) -> None:
    """
    Run the complete transcription and post-processing pipeline.

    This function automatically detects the audio language and:
        1. Transcribes using faster-whisper (large-v3 model)
        2. Exports the raw transcript with timestamps to a .docx file
        3. Cleans the transcript using a language-appropriate LLM
        4. Summarizes the transcript using an LLM
        5. Exports cleaned and summarized versions to separate .docx files

    Args:
        audio_file: Path to the audio file to process.

    Raises:
        FileNotFoundError: If the audio file doesn't exist.
        ValueError: If the audio format is not supported.
    """
    # Setup paths
    audio_path = Path(audio_file)
    output_dir = audio_path.parent
    base_name = audio_path.stem  # Filename without extension

    print("=" * 60)
    print("Multi-Language Audio Transcription Pipeline")
    print("=" * 60)

    # -------------------------------------------------------------------------
    # Step 1: Transcribe the audio file
    # -------------------------------------------------------------------------
    # The Transcriber uses faster-whisper with automatic language detection.
    # Language is detected from the first 30 seconds of audio.
    # 'large-v3' provides the best accuracy across all languages.

    print(f"\n[Step 1/4] Transcribing audio...")
    print(f"  Input file: {audio_path}")
    print(f"  Detecting language automatically...")

    transcriber = Transcriber(model_size="large-v3")
    result = transcriber.transcribe(audio_path)

    # Display transcription metadata
    detected_language = get_language_name(result.language)
    print(f"\n  Transcription complete!")
    print(f"  Detected language: {detected_language} ({result.language})")
    print(f"  Confidence: {result.language_probability:.1%}")
    print(f"  Duration: {format_duration(result.duration)}")
    print(f"  Segments: {len(result.segments)}")

    # -------------------------------------------------------------------------
    # Step 2: Export raw transcript to Word document
    # -------------------------------------------------------------------------
    # The raw transcript includes timestamps for each segment, which is useful
    # for referencing specific parts of the audio later.

    print(f"\n[Step 2/4] Saving raw transcript...")

    raw_docx_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_docx_path)

    print(f"  Saved to: {raw_docx_path}")

    # -------------------------------------------------------------------------
    # Step 3: Clean transcript with LLM
    # -------------------------------------------------------------------------
    # The PostProcessor.clean() method automatically:
    #   - Detects the primary language of the text
    #   - Selects the appropriate LLM model (see language_models mapping)
    #   - Corrects transcription errors
    #   - Removes filler words and repeated phrases
    #   - Adds proper punctuation

    print(f"\n[Step 3/4] Cleaning transcript with LLM...")

    processor = PostProcessor()

    # Display which model will be used for this language
    model_for_lang = processor.language_models.get(
        result.language,
        processor.language_models.get("other", processor.model),
    )
    print(f"  Using model: {model_for_lang} (for {detected_language})")

    cleaned_result = processor.clean(
        result.text,
        language=result.language,  # Pass detected language to skip re-detection
        chunk_size=500,  # Process in 500-character chunks
        show_progress=True,  # Show chunk progress
    )

    # Save cleaned transcript
    cleaned_docx_path = output_dir / f"{base_name}_cleaned.docx"
    save_text_to_docx(
        text=cleaned_result.text,
        output_path=cleaned_docx_path,
        title="Cleaned Transcript",
        language=result.language,
        duration=result.duration,
    )

    print(f"  Saved to: {cleaned_docx_path}")

    # -------------------------------------------------------------------------
    # Step 4: Summarize transcript with LLM
    # -------------------------------------------------------------------------
    # The summarize() method is language-agnostic and works with any language.
    # Available styles:
    #   - 'concise': Brief bullet points of key topics
    #   - 'detailed': Comprehensive summary with main points and conclusions
    #   - 'action_items': Extracts tasks, decisions, and responsibilities

    print(f"\n[Step 4/4] Summarizing transcript with LLM...")

    summary_result = processor.summarize(
        result.text,
        style="detailed",  # Options: 'concise', 'detailed', 'action_items'
        chunk_size=500,  # Process in 500-character chunks
        show_progress=True,  # Show chunk progress
    )

    # Save summary
    summary_docx_path = output_dir / f"{base_name}_summary.docx"
    save_text_to_docx(
        text=summary_result.text,
        output_path=summary_docx_path,
        title="Transcript Summary",
        language=result.language,
        duration=result.duration,
    )

    print(f"  Saved to: {summary_docx_path}")

    # -------------------------------------------------------------------------
    # Summary of outputs
    # -------------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("Processing Complete!")
    print("=" * 60)
    print(f"\nDetected Language: {detected_language} ({result.language})")
    print(f"Model Used: {model_for_lang}")
    print(f"\nOutput files created:")
    print(f"  1. Raw transcript:     {raw_docx_path}")
    print(f"  2. Cleaned transcript: {cleaned_docx_path}")
    print(f"  3. Summary:            {summary_docx_path}")
    print()


# =============================================================================
# Script Entry Point
# =============================================================================


if __name__ == "__main__":
    # Check command line arguments
    if len(sys.argv) != 2:
        print("Usage: python transcribe_auto_detect.py <audio_file>")
        print()
        print("This script automatically detects the language of your audio.")
        print()
        print("Example:")
        print("  python transcribe_auto_detect.py meeting.wav")
        print()
        print("Supported formats: .mp3, .wav, .m4a, .flac, .ogg, .opus, .wma, .aac")
        print()
        print("Supported languages: English, Chinese, Japanese, Korean, Spanish,")
        print("                     French, German, and 90+ others via Whisper")
        sys.exit(1)

    audio_file = sys.argv[1]

    # Verify file exists before processing
    if not Path(audio_file).exists():
        print(f"Error: File not found: {audio_file}")
        sys.exit(1)

    # Run the pipeline
    try:
        process_audio(audio_file)
    except KeyboardInterrupt:
        print("\n\nProcessing interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)
