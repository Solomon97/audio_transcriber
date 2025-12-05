"""
simple.py — Minimal script for audio transcription, cleaning, and summarization.

Edit the variables below, then run: python simple.py
"""

from pathlib import Path

# =============================================================================
# USER SETTINGS — Edit these variables
# =============================================================================

# Path to your audio file
AUDIO_FILE = "testing_bits/english_audio/test_folder/TED_Talks_Daily_Trailer.mp3"

# Output directory (same as audio file if empty)
OUTPUT_DIR = ""

# Whisper model size: "tiny", "base", "small", "medium", "large-v2", "large-v3"
# Larger = more accurate but slower. "large-v3" recommended for Chinese.
MODEL_SIZE = "large-v3"

# Language hint: "zh" for Chinese, "en" for English, or None for auto-detect
LANGUAGE = None

# LLM model for post-processing (requires Ollama running locally)
LLM_MODEL = "qwen2.5:7b"

# Output in Traditional Chinese? Set False for Simplified.
USE_TRADITIONAL = True

# Translation target language (e.g., "en", "zh-TW", "ja", or None to skip)
TRANSLATE_TO = "zh-TW"

# =============================================================================
# SCRIPT — No need to edit below this line
# =============================================================================

from docx import Document

from audio_transcriber import Transcriber, PostProcessor


def save_to_docx(text: str, path: Path, title: str) -> None:
    """Save text content to a Word document."""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    doc.save(path)


def main():
    # Setup paths
    audio_path = Path(AUDIO_FILE)
    output_dir = Path(OUTPUT_DIR) if OUTPUT_DIR else audio_path.parent
    base_name = audio_path.stem

    # --- Transcribe ---
    print(f"Transcribing: {audio_path}")
    transcriber = Transcriber(model_size=MODEL_SIZE)
    result = transcriber.transcribe(audio_path, language=LANGUAGE)

    print(f"Language: {result.language} ({result.language_probability:.0%})")
    print(f"Duration: {result.duration:.1f}s")

    # Save raw transcript
    raw_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_path)
    print(f"Saved: {raw_path}")

    # --- Clean with LLM ---
    print("\nCleaning...")
    processor = PostProcessor(model=LLM_MODEL)
    cleaned = processor.clean(result.text)

    cleaned_path = output_dir / f"{base_name}_cleaned.docx"
    save_to_docx(cleaned.text, cleaned_path, "Cleaned Transcript")
    print(f"Saved: {cleaned_path}")

    # --- Summarize with LLM ---
    print("\nSummarizing...")
    summary = processor.summarize_chinese(
        result.text,
        use_traditional=USE_TRADITIONAL,
        style="concise",
    )

    summary_path = output_dir / f"{base_name}_summary.docx"
    save_to_docx(summary.text, summary_path, "Summary")
    print(f"Saved: {summary_path}")

    # --- Translate with LLM ---
    if TRANSLATE_TO:
        print(f"\nTranslating to {TRANSLATE_TO}...")
        translated = processor.translate(result.text, target_language=TRANSLATE_TO)

        translated_path = output_dir / f"{base_name}_translated_{TRANSLATE_TO}.docx"
        save_to_docx(translated.text, translated_path, f"Translation ({TRANSLATE_TO})")
        print(f"Saved: {translated_path}")

    print("\nDone!")


if __name__ == "__main__":
    main()
