"""
simple.py — 音訊轉錄、清理、摘要的簡易腳本

修改下方變數後執行: python simple.py
"""

from pathlib import Path

# =============================================================================
# 使用者設定 — 請修改以下變數
# =============================================================================

# 音訊檔案路徑
AUDIO_FILE = "testing_bits/english_audio/test_folder/TED_Talks_Daily_Trailer.mp3"

# 輸出目錄（留空則與音訊檔案同目錄）
OUTPUT_DIR = ""

# Whisper 模型大小: "tiny", "base", "small", "medium", "large-v2", "large-v3"
# 越大越準確但越慢，中文建議使用 "large-v3"
MODEL_SIZE = "large-v3"

# 語言提示: "zh" 中文, "en" 英文, 或 None 自動偵測
LANGUAGE = None

# LLM 模型（需要本地運行 Ollama）
LLM_MODEL = "qwen2.5:7b"

# 輸出繁體中文？設為 False 則輸出簡體
USE_TRADITIONAL = True

# 翻譯目標語言（例如 "en", "zh-TW", "ja"，設為 None 則跳過翻譯）
TRANSLATE_TO = "zh-TW"

# =============================================================================
# 腳本內容 — 以下無需修改
# =============================================================================

from docx import Document

from audio_transcriber import Transcriber, PostProcessor


def save_to_docx(text: str, path: Path, title: str) -> None:
    """將文字內容儲存為 Word 文件"""
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(text)
    doc.save(path)


def main():
    # 設定路徑
    audio_path = Path(AUDIO_FILE)
    output_dir = Path(OUTPUT_DIR) if OUTPUT_DIR else audio_path.parent
    base_name = audio_path.stem

    # --- 轉錄 ---
    print(f"轉錄中: {audio_path}")
    transcriber = Transcriber(model_size=MODEL_SIZE)
    result = transcriber.transcribe(audio_path, language=LANGUAGE)

    print(f"語言: {result.language} ({result.language_probability:.0%})")
    print(f"時長: {result.duration:.1f} 秒")

    # 儲存原始逐字稿
    raw_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_path)
    print(f"已儲存: {raw_path}")

    # --- 清理 ---
    print("\n清理中...")
    processor = PostProcessor(model=LLM_MODEL)
    cleaned = processor.clean(result.text)

    cleaned_path = output_dir / f"{base_name}_cleaned.docx"
    save_to_docx(cleaned.text, cleaned_path, "清理後逐字稿")
    print(f"已儲存: {cleaned_path}")

    # --- 摘要 ---
    print("\n摘要中...")
    summary = processor.summarize_chinese(
        result.text,
        use_traditional=USE_TRADITIONAL,
        style="concise",
    )

    summary_path = output_dir / f"{base_name}_summary.docx"
    save_to_docx(summary.text, summary_path, "摘要")
    print(f"已儲存: {summary_path}")

    # --- 翻譯 ---
    if TRANSLATE_TO:
        print(f"\n翻譯至 {TRANSLATE_TO} 中...")
        translated = processor.translate(result.text, target_language=TRANSLATE_TO)

        translated_path = output_dir / f"{base_name}_translated_{TRANSLATE_TO}.docx"
        save_to_docx(translated.text, translated_path, f"翻譯 ({TRANSLATE_TO})")
        print(f"已儲存: {translated_path}")

    print("\n完成!")


if __name__ == "__main__":
    main()
