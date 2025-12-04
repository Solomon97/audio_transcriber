#!/usr/bin/env python3
"""
transcribe_and_extract.py — Example script for transcription with character extraction.

This script automatically detects the language of your audio, transcribes it,
and extracts the main characters/speakers mentioned in the transcript.

Workflow:
    1. Transcribe audio using faster-whisper (auto-detects language)
    2. Export the raw transcript to a Word document (.docx)
    3. Extract main characters using an LLM
    4. Display extracted characters and export to Word document

Prerequisites:
    - Install the package: pip install -e .[ollama]
    - Have Ollama running locally (default: http://localhost:11434)
    - Pull a capable model: ollama pull qwen2.5:7b

Usage:
    python transcribe_and_extract.py <audio_file>

Examples:
    python transcribe_and_extract.py interview.wav
    python transcribe_and_extract.py podcast.mp3
    python transcribe_and_extract.py meeting.m4a

Output:
    Creates two .docx files in the same directory as the input audio:
        - <filename>_raw.docx         : Raw transcript with timestamps
        - <filename>_characters.docx  : Extracted character information

What Gets Extracted:
    The LLM analyzes the transcript to identify:
        - Names of people mentioned or speaking
        - Their roles (interviewer, guest, manager, etc.)
        - Relationships between characters
        - Key details (occupation, affiliation, etc.)
        - Notable quotes or statements attributed to each person

Notes:
    - For best accuracy, the script uses the 'large-v3' Whisper model
    - Extraction quality depends on how clearly people are identified in the audio
    - Works best with interviews, podcasts, meetings, and conversations
"""

import sys
from pathlib import Path

from docx import Document

from audio_transcriber import Transcriber, PostProcessor


# =============================================================================
# Configuration
# =============================================================================

# Human-readable language names for display
LANGUAGE_NAMES: dict[str, str] = {
    "en": "English",
    "zh": "Chinese",
    "zh-TW": "Traditional Chinese",
    "zh-CN": "Simplified Chinese",
    "ja": "Japanese",
    "ko": "Korean",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "ru": "Russian",
    "pt": "Portuguese",
    "it": "Italian",
}

# Fields to extract from the transcript
# These are passed to processor.extract() to guide the LLM
EXTRACTION_FIELDS: list[str] = [
    "names of all people mentioned or speaking",
    "role or title of each person",
    "relationships between people",
    "key details about each person (occupation, affiliation, background)",
    "notable statements or quotes attributed to each person",
]


# =============================================================================
# Helper Functions
# =============================================================================


def get_language_name(code: str) -> str:
    """
    Get the human-readable name for a language code.

    Args:
        code: Language code (e.g., 'zh', 'en', 'zh-TW').

    Returns:
        Human-readable language name, or the code itself if unknown.
    """
    return LANGUAGE_NAMES.get(code, code.upper())


def save_text_to_docx(
    text: str,
    output_path: Path,
    title: str,
    language: str,
    duration: float,
) -> Path:
    """
    Save text content to a Word document with metadata.

    Args:
        text: The text content to save.
        output_path: Path where the .docx file will be saved.
        title: Document title/heading.
        language: Detected language code (e.g., 'zh', 'en').
        duration: Audio duration in seconds.

    Returns:
        Path to the created document.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    # Add metadata section
    language_display = f"{get_language_name(language)} ({language})"
    doc.add_paragraph(f"Language: {language_display}")
    doc.add_paragraph(f"Duration: {format_duration(duration)}")
    doc.add_paragraph()  # Empty line for spacing

    # Add main content
    doc.add_heading("Content", level=2)
    doc.add_paragraph(text)

    doc.save(output_path)
    return output_path


def save_extraction_to_docx(
    extraction_text: str,
    output_path: Path,
    language: str,
    duration: float,
) -> Path:
    """
    Save character extraction results to a Word document.

    Args:
        extraction_text: The extracted character information from LLM.
        output_path: Path where the .docx file will be saved.
        language: Detected language code (e.g., 'zh', 'en').
        duration: Audio duration in seconds.

    Returns:
        Path to the created document.
    """
    doc = Document()
    doc.add_heading("Character Extraction", level=1)

    # Add metadata section
    language_display = f"{get_language_name(language)} ({language})"
    doc.add_paragraph(f"Source Language: {language_display}")
    doc.add_paragraph(f"Audio Duration: {format_duration(duration)}")
    doc.add_paragraph()  # Empty line for spacing

    # Add extraction results
    doc.add_heading("Main Characters", level=2)
    doc.add_paragraph(extraction_text)

    doc.save(output_path)
    return output_path


def format_duration(seconds: float) -> str:
    """
    Format seconds into a human-readable duration string.

    Args:
        seconds: Duration in seconds.

    Returns:
        Formatted string like "5m 30s" or "1h 15m 30s".
    """
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)

    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    elif minutes > 0:
        return f"{minutes}m {secs}s"
    else:
        return f"{secs}s"


# =============================================================================
# Main Processing Pipeline
# =============================================================================


def process_audio(audio_file: str) -> None:
    """
    Run the transcription and character extraction pipeline.

    This function:
        1. Transcribes audio using faster-whisper (auto-detects language)
        2. Exports the raw transcript with timestamps to a .docx file
        3. Extracts main characters/speakers using an LLM
        4. Displays results and exports to a .docx file

    Args:
        audio_file: Path to the audio file to process.

    Raises:
        FileNotFoundError: If the audio file doesn't exist.
        ValueError: If the audio format is not supported.
    """
    # Setup paths
    audio_path = Path(audio_file)
    output_dir = audio_path.parent
    base_name = audio_path.stem  # Filename without extension

    print("=" * 60)
    print("Transcription and Character Extraction Pipeline")
    print("=" * 60)

    # -------------------------------------------------------------------------
    # Step 1: Transcribe the audio file
    # -------------------------------------------------------------------------
    # The Transcriber uses faster-whisper with automatic language detection.
    # Language is detected from the first 30 seconds of audio.

    print(f"\n[Step 1/3] Transcribing audio...")
    print(f"  Input file: {audio_path}")
    print(f"  Detecting language automatically...")

    transcriber = Transcriber(model_size="large-v3")
    result = transcriber.transcribe(audio_path)

    # Display transcription metadata
    source_language = result.language
    source_name = get_language_name(source_language)

    print(f"\n  Transcription complete!")
    print(f"  Detected language: {source_name} ({source_language})")
    print(f"  Confidence: {result.language_probability:.1%}")
    print(f"  Duration: {format_duration(result.duration)}")
    print(f"  Segments: {len(result.segments)}")

    # -------------------------------------------------------------------------
    # Step 2: Export raw transcript to Word document
    # -------------------------------------------------------------------------
    # The raw transcript includes timestamps for each segment.

    print(f"\n[Step 2/3] Saving raw transcript...")

    raw_docx_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_docx_path)

    print(f"  Saved to: {raw_docx_path}")

    # -------------------------------------------------------------------------
    # Step 3: Extract main characters
    # -------------------------------------------------------------------------
    # The PostProcessor.extract() method uses an LLM to identify and describe
    # the main people mentioned or speaking in the transcript.

    print(f"\n[Step 3/3] Extracting main characters...")
    print(f"  Analyzing transcript for people, roles, and relationships...")

    processor = PostProcessor()

    extraction_result = processor.extract(
        result.text,
        fields=EXTRACTION_FIELDS,
    )

    # Display extracted information in terminal
    print(f"\n" + "-" * 60)
    print("Extracted Characters:")
    print("-" * 60)
    print(extraction_result.text)
    print("-" * 60)

    # Save extraction to Word document
    characters_docx_path = output_dir / f"{base_name}_characters.docx"
    save_extraction_to_docx(
        extraction_text=extraction_result.text,
        output_path=characters_docx_path,
        language=source_language,
        duration=result.duration,
    )

    print(f"\n  Saved to: {characters_docx_path}")

    # -------------------------------------------------------------------------
    # Summary of outputs
    # -------------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("Processing Complete!")
    print("=" * 60)
    print(f"\nSource Language: {source_name} ({source_language})")
    print(f"\nOutput files created:")
    print(f"  1. Raw transcript:        {raw_docx_path}")
    print(f"  2. Character extraction:  {characters_docx_path}")
    print()


# =============================================================================
# Script Entry Point
# =============================================================================


if __name__ == "__main__":
    # Check command line arguments
    if len(sys.argv) != 2:
        print("Usage: python transcribe_and_extract.py <audio_file>")
        print()
        print("This script transcribes audio and extracts main characters/speakers.")
        print()
        print("Example:")
        print("  python transcribe_and_extract.py interview.wav")
        print()
        print("Supported formats: .mp3, .wav, .m4a, .flac, .ogg, .opus, .wma, .aac")
        print()
        print("What gets extracted:")
        print("  - Names of people mentioned or speaking")
        print("  - Their roles (interviewer, guest, manager, etc.)")
        print("  - Relationships between characters")
        print("  - Key details (occupation, affiliation, background)")
        print("  - Notable quotes attributed to each person")
        sys.exit(1)

    audio_file = sys.argv[1]

    # Verify file exists before processing
    if not Path(audio_file).exists():
        print(f"Error: File not found: {audio_file}")
        sys.exit(1)

    # Run the pipeline
    try:
        process_audio(audio_file)
    except KeyboardInterrupt:
        print("\n\nProcessing interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)
