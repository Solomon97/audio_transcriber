#!/usr/bin/env python3
"""
transcribe_and_translate.py — Example script for transcription with translation.

This script automatically detects the language of your audio, transcribes it,
and translates the transcript to your desired target language.

Workflow:
    1. Transcribe audio using faster-whisper (auto-detects language)
    2. Export the raw transcript to a Word document (.docx)
    3. Translate the transcript to the target language using an LLM
    4. Export the translated version to a Word document

Prerequisites:
    - Install the package: pip install -e .[ollama]
    - Have Ollama running locally (default: http://localhost:11434)
    - Pull a capable model: ollama pull qwen2.5:7b

Usage:
    python transcribe_and_translate.py <audio_file> [--to <language>]

Arguments:
    audio_file    Path to the audio file to transcribe
    --to          Target language code (default: en)

Examples:
    # Translate to English (default)
    python transcribe_and_translate.py meeting.wav

    # Translate Chinese audio to English
    python transcribe_and_translate.py chinese_meeting.mp3 --to en

    # Translate English audio to Traditional Chinese
    python transcribe_and_translate.py english_podcast.wav --to zh-TW

    # Translate to Japanese
    python transcribe_and_translate.py lecture.m4a --to ja

Output:
    Creates two .docx files in the same directory as the input audio:
        - <filename>_raw.docx         : Raw transcript with timestamps
        - <filename>_<lang>.docx      : Translated transcript (e.g., _en.docx)

Supported Target Languages:
    en      English
    zh-TW   Traditional Chinese
    zh-CN   Simplified Chinese
    zh      Chinese (general)
    ja      Japanese
    ko      Korean
    es      Spanish
    fr      French
    de      German
    ru      Russian
    pt      Portuguese
    it      Italian

Notes:
    - For best accuracy, the script uses the 'large-v3' Whisper model
    - Translation quality depends on the Ollama model used
    - Long transcripts are processed in chunks for better results
"""

import argparse
import sys
from pathlib import Path

from docx import Document

from audio_transcriber import Transcriber, PostProcessor


# =============================================================================
# Configuration
# =============================================================================

# Human-readable language names for display
LANGUAGE_NAMES: dict[str, str] = {
    "en": "English",
    "zh": "Chinese",
    "zh-TW": "Traditional Chinese",
    "zh-CN": "Simplified Chinese",
    "ja": "Japanese",
    "ko": "Korean",
    "es": "Spanish",
    "fr": "French",
    "de": "German",
    "ru": "Russian",
    "pt": "Portuguese",
    "it": "Italian",
    "nl": "Dutch",
    "pl": "Polish",
    "vi": "Vietnamese",
    "th": "Thai",
    "id": "Indonesian",
    "ar": "Arabic",
    "hi": "Hindi",
}

# Default target language for translation
DEFAULT_TARGET_LANGUAGE = "en"


# =============================================================================
# Helper Functions
# =============================================================================


def get_language_name(code: str) -> str:
    """
    Get the human-readable name for a language code.

    Args:
        code: Language code (e.g., 'zh', 'en', 'zh-TW').

    Returns:
        Human-readable language name, or the code itself if unknown.
    """
    return LANGUAGE_NAMES.get(code, code.upper())


def save_text_to_docx(
    text: str,
    output_path: Path,
    title: str,
    source_language: str,
    target_language: str | None,
    duration: float,
) -> Path:
    """
    Save text content to a Word document with metadata.

    Args:
        text: The text content to save.
        output_path: Path where the .docx file will be saved.
        title: Document title/heading.
        source_language: Original language code (e.g., 'zh', 'en').
        target_language: Target language code if translated, None for raw.
        duration: Audio duration in seconds.

    Returns:
        Path to the created document.
    """
    doc = Document()
    doc.add_heading(title, level=1)

    # Add metadata section
    source_display = f"{get_language_name(source_language)} ({source_language})"
    doc.add_paragraph(f"Source Language: {source_display}")

    if target_language:
        target_display = f"{get_language_name(target_language)} ({target_language})"
        doc.add_paragraph(f"Target Language: {target_display}")

    doc.add_paragraph(f"Duration: {format_duration(duration)}")
    doc.add_paragraph()  # Empty line for spacing

    # Add main content
    doc.add_heading("Content", level=2)
    doc.add_paragraph(text)

    doc.save(output_path)
    return output_path


def format_duration(seconds: float) -> str:
    """
    Format seconds into a human-readable duration string.

    Args:
        seconds: Duration in seconds.

    Returns:
        Formatted string like "5m 30s" or "1h 15m 30s".
    """
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)

    if hours > 0:
        return f"{hours}h {minutes}m {secs}s"
    elif minutes > 0:
        return f"{minutes}m {secs}s"
    else:
        return f"{secs}s"


def languages_match(source: str, target: str) -> bool:
    """
    Check if source and target languages are effectively the same.

    Handles cases like 'zh' matching 'zh-TW' or 'zh-CN'.

    Args:
        source: Source language code from Whisper.
        target: Target language code specified by user.

    Returns:
        True if languages are the same or closely related.
    """
    # Exact match
    if source == target:
        return True

    # Chinese variants: 'zh' from Whisper matches 'zh-TW', 'zh-CN'
    if source == "zh" and target in ("zh", "zh-TW", "zh-CN"):
        return True

    return False


def parse_args() -> argparse.Namespace:
    """
    Parse command line arguments.

    Returns:
        Parsed arguments namespace.
    """
    parser = argparse.ArgumentParser(
        description="Transcribe audio and translate to another language.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
            Supported target languages:
            en      English           zh-TW   Traditional Chinese
            zh-CN   Simplified Chinese    ja      Japanese
            ko      Korean            es      Spanish
            fr      French            de      German
            ru      Russian           pt      Portuguese
            it      Italian

            Examples:
            python transcribe_and_translate.py meeting.wav
            python transcribe_and_translate.py meeting.wav --to zh-TW
            python transcribe_and_translate.py lecture.mp3 --to ja
        """,
    )

    parser.add_argument(
        "audio_file",
        type=str,
        help="Path to the audio file to transcribe",
    )

    parser.add_argument(
        "--to",
        type=str,
        default=DEFAULT_TARGET_LANGUAGE,
        dest="target_language",
        metavar="LANG",
        help=f"Target language code (default: {DEFAULT_TARGET_LANGUAGE})",
    )

    return parser.parse_args()


# =============================================================================
# Main Processing Pipeline
# =============================================================================


def process_audio(audio_file: str, target_language: str) -> None:
    """
    Run the transcription and translation pipeline.

    This function:
        1. Transcribes audio using faster-whisper (auto-detects language)
        2. Exports the raw transcript with timestamps to a .docx file
        3. Translates the transcript to the target language
        4. Exports the translated version to a .docx file

    Args:
        audio_file: Path to the audio file to process.
        target_language: Target language code for translation.

    Raises:
        FileNotFoundError: If the audio file doesn't exist.
        ValueError: If the audio format is not supported.
    """
    # Setup paths
    audio_path = Path(audio_file)
    output_dir = audio_path.parent
    base_name = audio_path.stem  # Filename without extension

    print("=" * 60)
    print("Transcription and Translation Pipeline")
    print("=" * 60)
    print(f"Target language: {get_language_name(target_language)} ({target_language})")

    # -------------------------------------------------------------------------
    # Step 1: Transcribe the audio file
    # -------------------------------------------------------------------------
    # The Transcriber uses faster-whisper with automatic language detection.
    # Language is detected from the first 30 seconds of audio.

    print(f"\n[Step 1/3] Transcribing audio...")
    print(f"  Input file: {audio_path}")
    print(f"  Detecting language automatically...")

    transcriber = Transcriber(model_size="large-v3")
    result = transcriber.transcribe(audio_path)

    # Display transcription metadata
    source_language = result.language
    source_name = get_language_name(source_language)

    print(f"\n  Transcription complete!")
    print(f"  Detected language: {source_name} ({source_language})")
    print(f"  Confidence: {result.language_probability:.1%}")
    print(f"  Duration: {format_duration(result.duration)}")
    print(f"  Segments: {len(result.segments)}")

    # -------------------------------------------------------------------------
    # Step 2: Export raw transcript to Word document
    # -------------------------------------------------------------------------
    # The raw transcript includes timestamps for each segment.

    print(f"\n[Step 2/3] Saving raw transcript...")

    raw_docx_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_docx_path)

    print(f"  Saved to: {raw_docx_path}")

    # -------------------------------------------------------------------------
    # Step 3: Translate transcript
    # -------------------------------------------------------------------------
    # Check if translation is needed (source != target)

    print(f"\n[Step 3/3] Translating transcript...")
    print(
        f"  {source_name} ({source_language}) → {get_language_name(target_language)} ({target_language})"
    )

    if languages_match(source_language, target_language):
        print(f"\n  Skipping translation: source and target languages are the same.")
        print(
            f"  The raw transcript is already in {get_language_name(target_language)}."
        )

        # Still create the "translated" file as a copy for consistency
        translated_docx_path = output_dir / f"{base_name}_{target_language}.docx"
        save_text_to_docx(
            text=result.text,
            output_path=translated_docx_path,
            title=f"Transcript ({get_language_name(target_language)})",
            source_language=source_language,
            target_language=None,  # No translation performed
            duration=result.duration,
        )
        print(f"  Saved to: {translated_docx_path}")

    else:
        # Perform translation
        processor = PostProcessor()

        translated_result = processor.translate(
            result.text,
            target_language=target_language,
            chunk_size=500,  # Process in 500-character chunks
            show_progress=True,  # Show chunk progress
        )

        # Save translated transcript
        translated_docx_path = output_dir / f"{base_name}_{target_language}.docx"
        save_text_to_docx(
            text=translated_result.text,
            output_path=translated_docx_path,
            title=f"Translated Transcript ({get_language_name(target_language)})",
            source_language=source_language,
            target_language=target_language,
            duration=result.duration,
        )

        print(f"  Saved to: {translated_docx_path}")

    # -------------------------------------------------------------------------
    # Summary of outputs
    # -------------------------------------------------------------------------
    print("\n" + "=" * 60)
    print("Processing Complete!")
    print("=" * 60)
    print(f"\nSource: {source_name} ({source_language})")
    print(f"Target: {get_language_name(target_language)} ({target_language})")
    print(f"\nOutput files created:")
    print(f"  1. Raw transcript:        {raw_docx_path}")
    print(f"  2. Translated transcript: {translated_docx_path}")
    print()


# =============================================================================
# Script Entry Point
# =============================================================================


if __name__ == "__main__":
    args = parse_args()

    audio_file = args.audio_file
    target_language = args.target_language

    # Verify file exists before processing
    if not Path(audio_file).exists():
        print(f"Error: File not found: {audio_file}")
        sys.exit(1)

    # Validate target language (warn if unknown, but allow it)
    if target_language not in LANGUAGE_NAMES:
        print(f"Warning: Unknown language code '{target_language}'.")
        print(f"         Translation will still be attempted.")
        print()

    # Run the pipeline
    try:
        process_audio(audio_file, target_language)
    except KeyboardInterrupt:
        print("\n\nProcessing interrupted by user.")
        sys.exit(1)
    except Exception as e:
        print(f"\nError: {e}")
        sys.exit(1)
