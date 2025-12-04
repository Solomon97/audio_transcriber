#!/usr/bin/env python3
"""
test_transcriber.py — Simple test script for the audio_transcriber package.

This script demonstrates the core functionality of the audio_transcriber package:
    1. Transcribe a Chinese audio recording
    2. Detect and display the language of the transcript
    3. Export the raw transcript to a Word document
    4. Clean up the transcript using an LLM and export to Word
    5. Summarize the transcript using an LLM and export to Word

Usage:
    python test_transcriber.py <audio_file>

Example:
    python test_transcriber.py meeting.wav

Requirements:
    - audio_transcriber package installed
    - Ollama running locally (for LLM post-processing)
    - A Chinese audio file to transcribe
"""

import sys
from pathlib import Path

from audio_transcriber import Transcriber, PostProcessor


def main(audio_file: str) -> None:
    """
    Run the full transcription and post-processing pipeline.

    Args:
        audio_file: Path to the audio file to transcribe.
    """
    audio_path = Path(audio_file)
    output_dir = audio_path.parent
    base_name = audio_path.stem

    # -------------------------------------------------------------------------
    # Step 1: Transcribe the audio file
    # -------------------------------------------------------------------------
    print(f"Transcribing: {audio_path}")
    print("-" * 50)

    transcriber = Transcriber(model_size="large-v3")
    result = transcriber.transcribe(audio_path)

    # -------------------------------------------------------------------------
    # Step 2: Display detected language
    # -------------------------------------------------------------------------
    print(f"Detected language: {result.language}")
    print(f"Language confidence: {result.language_probability:.1%}")
    print(f"Audio duration: {result.duration:.1f} seconds")
    print("-" * 50)

    # -------------------------------------------------------------------------
    # Step 3: Export raw transcript to Word document
    # -------------------------------------------------------------------------
    raw_docx_path = output_dir / f"{base_name}_raw.docx"
    result.to_docx(raw_docx_path)
    print(f"Raw transcript saved to: {raw_docx_path}")

    # -------------------------------------------------------------------------
    # Step 4: Clean up transcript with LLM and export to Word
    # -------------------------------------------------------------------------
    print("\nCleaning transcript with LLM...")

    processor = PostProcessor()
    cleaned_result = processor.clean_chinese(result.text, use_traditional=True)

    # Save cleaned text to Word document
    cleaned_docx_path = output_dir / f"{base_name}_cleaned.docx"
    save_text_to_docx(
        text=cleaned_result.text,
        output_path=cleaned_docx_path,
        title="Cleaned Transcript",
        language=result.language,
        duration=result.duration,
    )
    print(f"Cleaned transcript saved to: {cleaned_docx_path}")

    # -------------------------------------------------------------------------
    # Step 5: Summarize transcript with LLM and export to Word
    # -------------------------------------------------------------------------
    print("\nSummarizing transcript with LLM...")

    summary_result = processor.summarize_chinese(
        result.text,
        use_traditional=True,
        style="detailed",
    )

    # Save summary to Word document
    summary_docx_path = output_dir / f"{base_name}_summary.docx"
    save_text_to_docx(
        text=summary_result.text,
        output_path=summary_docx_path,
        title="Transcript Summary",
        language=result.language,
        duration=result.duration,
    )
    print(f"Summary saved to: {summary_docx_path}")

    # -------------------------------------------------------------------------
    # Done
    # -------------------------------------------------------------------------
    print("\n" + "=" * 50)
    print("Processing complete!")
    print(f"  Raw transcript:     {raw_docx_path}")
    print(f"  Cleaned transcript: {cleaned_docx_path}")
    print(f"  Summary:            {summary_docx_path}")


def save_text_to_docx(
    text: str,
    output_path: Path,
    title: str,
    language: str,
    duration: float,
) -> None:
    """
    Save text content to a Word document with metadata.

    Args:
        text: The text content to save.
        output_path: Path where the .docx file will be saved.
        title: Document title/heading.
        language: Detected language code.
        duration: Audio duration in seconds.
    """
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)

    # Add metadata
    doc.add_paragraph(f"Language: {language}")
    doc.add_paragraph(f"Duration: {duration:.1f} seconds")
    doc.add_paragraph()

    # Add content
    doc.add_heading("Content", level=2)
    doc.add_paragraph(text)

    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        print(__doc__)
        sys.exit(1)

    main(sys.argv[1])
