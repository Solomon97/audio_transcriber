"""
transcriber.py — Core transcription functionality using faster-whisper.

This module provides the main Transcriber class for converting audio files
to text using the faster-whisper library, with automatic language detection
optimized for Chinese and English content.
"""

from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from faster_whisper import WhisperModel
from opencc import OpenCC

from .config import DEFAULT_MODEL_SIZE, SUPPORTED_AUDIO_FORMATS
from .utils import validate_audio_file, format_timestamp


ModelSize = Literal["tiny", "base", "small", "medium", "large-v2", "large-v3"]


@dataclass
class Segment:
    """
    A single transcription segment with timing information.

    Attributes:
        start: Start time in seconds.
        end: End time in seconds.
        text: Transcribed text for this segment.
    """

    start: float
    end: float
    text: str

    @property
    def duration(self) -> float:
        """Return the duration of this segment in seconds."""
        return self.end - self.start

    def __str__(self) -> str:
        """Return a formatted string with timestamps and text."""
        start_ts = format_timestamp(self.start)
        end_ts = format_timestamp(self.end)
        return f"[{start_ts} -> {end_ts}] {self.text}"


@dataclass
class TranscriptionResult:
    """
    Complete transcription result with full text and individual segments.

    Attributes:
        text: The complete transcribed text as a single string.
        segments: List of individual segments with timing information.
        language: Detected or specified language code (e.g., 'zh', 'en').
        language_probability: Confidence score for language detection (0.0-1.0).
        duration: Total audio duration in seconds.
    """

    text: str
    segments: list[Segment] = field(default_factory=list)
    language: str = ""
    language_probability: float = 0.0
    duration: float = 0.0

    def to_docx(self, output_path: str | Path) -> Path:
        """
        Export the transcription to a Word document.

        Args:
            output_path: Path where the .docx file will be saved.

        Returns:
            Path to the created document.

        Raises:
            ValueError: If output_path doesn't end with .docx extension.
        """
        from docx import Document

        output_path = Path(output_path)
        if output_path.suffix.lower() != ".docx":
            raise ValueError(f"Output path must have .docx extension: {output_path}")

        doc = Document()
        doc.add_heading("Transcription", level=1)

        # Add metadata
        doc.add_paragraph(f"Language: {self.language}")
        doc.add_paragraph(f"Duration: {format_timestamp(self.duration)}")
        doc.add_paragraph()

        # Add segments with timestamps
        doc.add_heading("Transcript", level=2)
        for segment in self.segments:
            doc.add_paragraph(str(segment))

        doc.save(output_path)
        return output_path

    def to_srt(self, output_path: str | Path) -> Path:
        """
        Export the transcription to SRT subtitle format.

        Args:
            output_path: Path where the .srt file will be saved.

        Returns:
            Path to the created file.
        """
        output_path = Path(output_path)
        lines = []

        for i, segment in enumerate(self.segments, start=1):
            start_ts = format_timestamp(segment.start, srt_format=True)
            end_ts = format_timestamp(segment.end, srt_format=True)
            lines.append(str(i))
            lines.append(f"{start_ts} --> {end_ts}")
            lines.append(segment.text.strip())
            lines.append("")

        output_path.write_text("\n".join(lines), encoding="utf-8")
        return output_path


class Transcriber:
    """
    Audio transcription engine using faster-whisper.

    This class provides a simple interface for transcribing audio files, with automatic language detection optimized for Chinese and English.

    Attributes:
        model_size: The Whisper model size being used.
        device: Compute device ('cpu', 'cuda', or 'auto').
        compute_type: Computation precision type.

    Example:
        >>> transcriber = Transcriber(model_size="large-v3")
        >>> result = transcriber.transcribe("meeting.wav")
        >>> print(result.text)
        >>> result.to_docx("meeting_transcript.docx")
    """

    def __init__(
        self,
        model_size: ModelSize = DEFAULT_MODEL_SIZE,
        device: str = "auto",
        compute_type: str = "default",
    ) -> None:
        """
        Initialize the Transcriber with a specified Whisper model.

        Args:
            model_size:
                Whisper model size. Options are 'tiny', 'base', 'small', 'medium', 'large-v2', or 'large-v3'.
                Larger models are more accurate but slower. Default is defined in config.py.
            device:
                Compute device to use. Options are 'cpu', 'cuda', or 'auto'. 'auto' will use CUDA if available, otherwise CPU.
            compute_type:
                Computation type for inference. Options include 'default', 'float16', 'int8_float16', 'int8'.
                Use 'int8' for lower memory usage on CPU.

        Raises:
            ValueError: If an invalid model_size is provided.
        """
        self.model_size = model_size
        self.device = device
        self.compute_type = compute_type
        self._model: WhisperModel | None = None

    @property
    def model(self) -> WhisperModel:
        """
        Lazy-load and return the Whisper model.

        The model is loaded on first access to avoid slow initialization
        when the Transcriber is created but not immediately used.

        Returns:
            The loaded WhisperModel instance.
        """
        if self._model is None:
            self._model = WhisperModel(
                self.model_size,
                device=self.device,
                compute_type=self.compute_type,
            )
        return self._model

    def transcribe(
        self,
        audio_path: str | Path,
        language: str | None = None,
        initial_prompt: str | None = None,
        word_timestamps: bool = False,
        vad_filter: bool = True,
    ) -> TranscriptionResult:
        """
        Transcribe an audio file to text.

        Args:
            audio_path: Path to the audio file. Supported formats are defined in config.SUPPORTED_AUDIO_FORMATS.
            language:
                Language code (e.g., 'zh', 'en').
                If None, language is automatically detected. For mixed Chinese/English content, setting language='zh' often works well.
            initial_prompt:
                Optional text to guide the transcription style. Useful for domain-specific vocabulary or formatting hints.
            word_timestamps:
                If True, include word-level timing information. This increases processing time.
            vad_filter:
                If True, use voice activity detection to filter out non-speech segments. Recommended for most use cases.

        Returns:
            TranscriptionResult containing the full text, segments with timestamps, and metadata.

        Raises:
            FileNotFoundError: If the audio file doesn't exist.
            ValueError: If the audio format is not supported.

        Example:
            >>> transcriber = Transcriber()
            >>> result = transcriber.transcribe("lecture.mp3", language="zh")
            >>> print(f"Detected language: {result.language}")
            >>> for segment in result.segments:
            ...     print(segment)
        """
        audio_path = Path(audio_path)
        validate_audio_file(audio_path)

        segments_iter, info = self.model.transcribe(
            str(audio_path),
            language=language,
            initial_prompt=initial_prompt,
            word_timestamps=word_timestamps,
            vad_filter=vad_filter,
        )
        # Convert all text to traditional Chinese since fasterwhisper tends to output simplified chinese
        converter = OpenCC("s2t")

        # Convert generator to list and build result
        segments = [
            Segment(start=seg.start, end=seg.end, text=converter.convert(seg.text))
            for seg in segments_iter
        ]

        full_text = " ".join(seg.text.strip() for seg in segments)

        return TranscriptionResult(
            text=full_text,
            segments=segments,
            language=info.language,
            language_probability=info.language_probability,
            duration=info.duration,
        )
